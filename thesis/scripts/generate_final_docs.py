import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_formatted_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(16 if level == 0 else 14 if level == 1 else 12)
        run.font.bold = True
    return h

def create_docx():
    doc = Document()
    
    # Global styles setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- Title Page ---
    add_formatted_heading(doc, 'State-Dependent Resistance to Ketamine-Augmented Psychotherapy: Baseline Network Instability as a Barrier to Recovery in PTSD', 0)
    
    doc.add_paragraph('\n' * 5)
    author = doc.add_paragraph('Author: [Your Name]')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: Sunday, June 14, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Institution: [Your Institution]').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # --- Abstract ---
    add_formatted_heading(doc, 'Abstract', 1)
    abstract_text = (
        "Background: Prolonged Exposure (PE) therapy is a highly effective treatment for Post-Traumatic Stress Disorder (PTSD), "
        "yet treatment resistance remains high. Emerging network models suggest that temporal instability within fronto-limbic fear circuits "
        "may hinder the consolidation of extinction learning. Subanesthetic ketamine is hypothesized to enhance neuroplasticity. "
        "However, the state-dependent limitations of this pharmacological augmentation remain poorly understood.\n\n"
        "Methods: In a double-blind RCT (N=30), PTSD patients underwent 1-week intensive PE augmented with high-dose ketamine, "
        "low-dose ketamine, or placebo. Baseline and post-treatment fMRI were analyzed for SN-DMN dynamic functional connectivity (dFC) variance.\n\n"
        "Results: In the High-Dose Ketamine group, baseline instability served as a robust prognostic barrier, strongly predicting "
        "treatment failure (r = -0.77, p = 0.005). Conversely, the Placebo group started significantly more stable, and their improvement "
        "was independent of baseline dynamics. Further, Ketamine-induced entropy spikes were only beneficial for placebo patients.\n\n"
        "Conclusions: Baseline network instability acts as a robust barrier to therapy success even with ketamine augmentation. "
        "Patients with highly erratic fear circuits possess neural noise that pharmacological plasticity alone cannot override."
    )
    doc.add_paragraph(abstract_text)
    doc.add_page_break()
    
    # --- Section 1: Introduction ---
    add_formatted_heading(doc, '1. Introduction', 1)
    doc.add_paragraph("PTSD is associated with dysregulation of large-scale brain networks, particularly the Salience Network (SN) and Default Mode Network (DMN).")
    
    p_fig1 = doc.add_paragraph()
    p_fig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fig1 = p_fig1.add_run()
    run_fig1.add_picture('thesis/paper/figures/high_res_brain_render.png', width=Inches(5))
    doc.add_paragraph("Figure 1: High-resolution surface rendering of target networks.")

    # --- Section 2: Methods ---
    add_formatted_heading(doc, '2. Materials and Methods', 1)
    doc.add_paragraph("Dynamic Functional Connectivity (dFC) was calculated for the SN-DMN axis.")

    # --- Section 3: Results ---
    add_formatted_heading(doc, '3. Results', 1)
    
    add_formatted_heading(doc, '3.1 Baseline Instability as a Prognostic Barrier', 2)
    doc.add_paragraph("In the high-dose ketamine arm, baseline SN-DMN instability strongly predicted clinical failure (r = -0.770, p = 0.0056).")
    
    p_fig5 = doc.add_paragraph()
    p_fig5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fig5 = p_fig5.add_run()
    run_fig5.add_picture('thesis/paper/figures/figure1_placebo.png', width=Inches(4.5))
    doc.add_paragraph("Figure 5: Baseline Instability Predicts Failure in Ketamine Group.")

    add_formatted_heading(doc, '3.2 Summary Finding', 2)
    p_fig4 = doc.add_paragraph()
    p_fig4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fig4 = p_fig4.add_run()
    run_fig4.add_picture('thesis/paper/figures/data_explanatory_summary.png', width=Inches(5.5))
    doc.add_paragraph("Figure 4: Integrated Summary of the Prognostic Barrier.")

    # --- Section 4: Discussion ---
    add_formatted_heading(doc, '4. Discussion', 1)
    doc.add_paragraph("Contrary to the rescue hypothesis, high-dose ketamine did not override pathological network states. Instead, severe baseline instability remains a definitive predictor of treatment failure.")

    p_fig12 = doc.add_paragraph()
    p_fig12.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fig12 = p_fig12.add_run()
    run_fig12.add_picture('thesis/paper/figures/conceptual_state_space.png', width=Inches(5))
    doc.add_paragraph("Figure 12: Neural State-Space Model of the Prognostic Barrier.")

    doc.save('thesis/manuscript.docx')
    print("DOCX generated: thesis/manuscript.docx")

if __name__ == "__main__":
    create_docx()
